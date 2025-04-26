# app/document_processing/resume_parser.py
import io
import re
import PyPDF2
import docx
from pdfminer.high_level import extract_text

def parse_resume(uploaded_file):
    """
    Parse resume file (PDF, DOCX, TXT) and extract relevant information
    """
    file_type = uploaded_file.name.split('.')[-1].lower()
    content = None
    
    # Extract text based on file type
    if file_type == 'pdf':
        content = extract_text_from_pdf(uploaded_file)
    elif file_type == 'docx':
        content = extract_text_from_docx(uploaded_file)
    elif file_type == 'txt':
        content = uploaded_file.getvalue().decode('utf-8')
    else:
        return None
    
    # Parse the content to extract structured information
    parsed_data = extract_resume_information(content)
    return parsed_data

def extract_text_from_pdf(pdf_file):
    """Extract text content from PDF file"""
    text = ""
    try:
        # First approach with PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_file.getvalue()))
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        # If PyPDF2 didn't extract text well, try pdfminer
        if not text.strip():
            text = extract_text(io.BytesIO(pdf_file.getvalue()))
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        # Fallback to pdfminer
        try:
            text = extract_text(io.BytesIO(pdf_file.getvalue()))
        except Exception as e2:
            print(f"Error with fallback PDF extraction: {e2}")
    
    return text

def extract_text_from_docx(docx_file):
    """Extract text content from DOCX file"""
    try:
        doc = docx.Document(io.BytesIO(docx_file.getvalue()))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

def extract_resume_information(text):
    """
    Extract structured information from the resume text
    """
    # Basic extraction - in a real app, use more sophisticated NLP
    resume_data = {
        'raw_text': text,
        'name': extract_name(text),
        'email': extract_email(text),
        'phone': extract_phone(text),
        'education': extract_education(text),
        'experience': extract_experience(text),
        'skills': extract_skills(text)
    }
    
    return resume_data

def extract_name(text):
    # Simple name extraction - first line of the resume often contains the name
    lines = text.split('\n')
    for line in lines[:5]:  # Check first 5 lines
        if line.strip() and not re.search(r'@|resume|cv|curriculum', line.lower()):
            return line.strip()
    return ""

def extract_email(text):
    # Email extraction using regex
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    emails = re.findall(email_pattern, text)
    return emails[0] if emails else ""

def extract_phone(text):
    # Phone number extraction using regex
    phone_pattern = r'(\+\d{1,2}\s?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}'
    phones = re.findall(phone_pattern, text)
    return phones[0] if phones else ""

def extract_education(text):
    # Basic education extraction
    edu_keywords = ['education', 'university', 'college', 'bachelor', 'master', 'phd', 'degree']
    education_section = extract_section(text, edu_keywords)
    
    # In a real app, parse this section further to extract institutions, degrees, years
    return education_section

def extract_experience(text):
    # Basic work experience extraction
    exp_keywords = ['experience', 'employment', 'work history', 'professional experience']
    experience_section = extract_section(text, exp_keywords)
    
    # In a real app, parse this further to extract companies, roles, dates, responsibilities
    return experience_section

def extract_skills(text):
    # Basic skills extraction
    skill_keywords = ['skills', 'technical skills', 'competencies', 'proficiencies']
    skills_section = extract_section(text, skill_keywords)
    
    # Extract individual skills
    # In a real app, use a predefined list of skills to match against
    return skills_section

def extract_section(text, keywords):
    """Extract a section from the text based on keywords"""
    text_lower = text.lower()
    for keyword in keywords:
        if keyword in text_lower:
            # Find the starting position of the keyword
            start_pos = text_lower.find(keyword)
            
            # Extract from this position to the next section
            section_text = text[start_pos:start_pos + 1000]  # Limit to 1000 chars for simplicity
            
            # In a real app, find the end of the section more intelligently
            return section_text
    
    return ""


# app/document_processing/jd_parser.py
import io
import re
import PyPDF2
import docx
from pdfminer.high_level import extract_text

def parse_jd_file(uploaded_file):
    """Parse job description from uploaded file"""
    file_type = uploaded_file.name.split('.')[-1].lower()
    content = None
    
    # Extract text based on file type
    if file_type == 'pdf':
        content = extract_text_from_pdf(uploaded_file)
    elif file_type == 'docx':
        content = extract_text_from_docx(uploaded_file)
    elif file_type == 'txt':
        content = uploaded_file.getvalue().decode('utf-8')
    else:
        return None
    
    # Parse the content to extract structured information
    return parse_jd_text(content)

def extract_text_from_pdf(pdf_file):
    """Extract text content from PDF file"""
    # Same implementation as in resume_parser.py
    text = ""
    try:
        # First approach with PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_file.getvalue()))
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        # If PyPDF2 didn't extract text well, try pdfminer
        if not text.strip():
            text = extract_text(io.BytesIO(pdf_file.getvalue()))
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        # Fallback to pdfminer
        try:
            text = extract_text(io.BytesIO(pdf_file.getvalue()))
        except Exception as e2:
            print(f"Error with fallback PDF extraction: {e2}")
    
    return text

def extract_text_from_docx(docx_file):
    """Extract text content from DOCX file"""
    # Same implementation as in resume_parser.py
    try:
        doc = docx.Document(io.BytesIO(docx_file.getvalue()))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

def parse_jd_text(text):
    """Parse job description from text"""
    jd_data = {
        'raw_text': text,
        'title': extract_job_title(text),
        'company': extract_company_name(text),
        'required_skills': extract_required_skills(text),
        'responsibilities': extract_responsibilities(text),
        'qualifications': extract_qualifications(text),
        'location': extract_location(text)
    }
    
    return jd_data

def extract_job_title(text):
    """Extract job title from JD text"""
    # Simple approach - look for common patterns
    title_patterns = [
        r'job title\s*[:-]\s*(.*?)[\n\r]',
        r'position\s*[:-]\s*(.*?)[\n\r]',
        r'role\s*[:-]\s*(.*?)[\n\r]'
    ]
    
    for pattern in title_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    
    # Fallback: First line is often the job title
    first_line = text.strip().split('\n')[0]
    if len(first_line) < 100:  # Reasonable length for a title
        return first_line
    
    return "Unknown Position"

def extract_company_name(text):
    """Extract company name from JD text"""
    company_patterns = [
        r'company\s*[:-]\s*(.*?)[\n\r]',
        r'about\s+([a-z0-9\s]+)[\n\r]',
        r'at\s+([a-z0-9\s]+)[\n\r]',
        r'join\s+([a-z0-9\s]+)[\n\r]'
    ]
    
    for pattern in company_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    
    return "Unknown Company"

def extract_required_skills(text):
    """Extract required skills from JD text"""
    # Look for skills section
    skills_section = extract_section(text, ['skills', 'requirements', 'qualifications'])
    
    # In a real app, use more sophisticated NLP to extract specific skills
    return skills_section

def extract_responsibilities(text):
    """Extract job responsibilities from JD text"""
    # Look for responsibilities section
    resp_section = extract_section(text, ['responsibilities', 'duties', 'what you\'ll do'])
    
    return resp_section

def extract_qualifications(text):
    """Extract required qualifications from JD text"""
    # Look for qualifications section
    qual_section = extract_section(text, ['qualifications', 'requirements', 'what we\'re looking for'])
    
    return qual_section

def extract_location(text):
    """Extract job location from JD text"""
    location_patterns = [
        r'location\s*[:-]\s*(.*?)[\n\r]',
        r'based in\s+(.*?)[\n\r]',
        r'position is (?:located|based) in\s+(.*?)[\n\r]'
    ]
    
    for pattern in location_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    
    return "Unknown Location"

def extract_section(text, keywords):
    """Extract a section from the text based on keywords"""
    text_lower = text.lower()
    for keyword in keywords:
        if keyword in text_lower:
            # Find the starting position of the keyword
            start_pos = text_lower.find(keyword)
            
            # Extract from this position to the next section
            section_text = text[start_pos:start_pos + 1000]  # Limit to 1000 chars for simplicity
            
            # In a real app, find the end of the section more intelligently
            return section_text
    
    return ""


# app/document_processing/web_scraper.py
import requests
from bs4 import BeautifulSoup
import re

def scrape_job_description(url):
    """Scrape job description from a URL"""
    try:
        # Get the webpage content
        response = requests.get(url, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
        response.raise_for_status()
        
        # Parse HTML
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Extract job information based on common patterns
        # This is simplified - in a real app, you would need site-specific extractors
        
        # Extract main content - this is highly site-dependent
        main_content = extract_main_content(soup)
        
        # Parse the extracted content
        from app.document_processing.jd_parser import parse_jd_text
        jd_data = parse_jd_text(main_content)
        
        # Add source URL
        jd_data['source_url'] = url
        
        return jd_data
    
    except Exception as e:
        print(f"Error scraping job description: {e}")
        return None

def extract_main_content(soup):
    """Extract the main content from the HTML"""
    # Try different common containers
    content = ""
    
    # Try article tag
    article = soup.find('article')
    if article:
        content = article.get_text()
    
    # Try job description specific divs
    if not content:
        job_divs = soup.find_all(['div', 'section'], class_=re.compile(r'job(-|\s)?desc|description|posting|details', re.I))
        if job_divs:
            content = job_divs[0].get_text()
    
    # Fallback to main content area
    if not content:
        main = soup.find(['main', 'div'], id=re.compile(r'main|content', re.I))
        if main:
            content = main.get_text()
    
    # Last resort - get all text
    if not content:
        content = soup.get_text()
    
    # Clean up the text
    content = re.sub(r'\s+', ' ', content).strip()
    
    return content


# app/interview/rag_engine.py
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.text_splitter import Rec
# 
# app/interview/rag_engine.py
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.llms import HuggingFaceHub
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
import os

class RAGEngine:
    def __init__(self, api_key=None):
        """Initialize the RAG engine with HuggingFace model"""
        # Set HuggingFace API key if provided
        if api_key:
            os.environ["HUGGINGFACE_API_KEY"] = api_key
        
        # Initialize embeddings model
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-mpnet-base-v2"
        )
        
        # Initialize LLM
        self.llm = HuggingFaceHub(
            repo_id="google/flan-t5-xl",
            model_kwargs={"temperature": 0.5, "max_length": 512}
        )
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        
        self.vectorstore = None
        self.qa_chain = None

    def create_knowledge_base(self, resume_data, jd_data):
        """Create knowledge base from resume and job description data"""
        # Combine resume and JD data
        documents = []
        
        # Add resume sections with metadata
        resume_sections = {
            "education": resume_data.get("education", ""),
            "experience": resume_data.get("experience", ""),
            "skills": resume_data.get("skills", "")
        }
        
        for section_name, content in resume_sections.items():
            if content:
                documents.append(f"RESUME {section_name.upper()}: {content}")
        
        # Add JD sections with metadata
        jd_sections = {
            "requirements": jd_data.get("required_skills", ""),
            "responsibilities": jd_data.get("responsibilities", ""),
            "qualifications": jd_data.get("qualifications", "")
        }
        
        for section_name, content in jd_sections.items():
            if content:
                documents.append(f"JOB DESCRIPTION {section_name.upper()}: {content}")
        
        # Split documents into chunks
        chunks = self.text_splitter.create_documents(documents)
        
        # Create vector store
        self.vectorstore = FAISS.from_documents(chunks, self.embeddings)
        
        # Create QA chain
        interview_prompt_template = """
        You are preparing interview questions for a job candidate.
        
        The candidate's resume information and job description are provided below.
        
        Based on this context, generate a relevant interview question.
        
        Context: {context}
        
        Question type: {question_type}
        Interview stage: {interview_stage}
        
        Generate a professional interview question:
        """
        
        PROMPT = PromptTemplate(
            template=interview_prompt_template,
            input_variables=["context", "question_type", "interview_stage"]
        )
        
        self.qa_chain = RetrievalQA.from_chain_type(
            llm=self.llm,
            chain_type="stuff",
            retriever=self.vectorstore.as_retriever(),
            chain_type_kwargs={"prompt": PROMPT}
        )
        
        return True

    def generate_question(self, question_type, interview_stage):
        """Generate an interview question based on the knowledge base"""
        if not self.qa_chain:
            raise ValueError("Knowledge base not initialized. Call create_knowledge_base first.")
        
        # Run the QA chain
        result = self.qa_chain({"question_type": question_type, "interview_stage": interview_stage})
        
        return result["result"]


# app/interview/question_generator.py
import random
from app.interview.rag_engine import RAGEngine

# Question types
QUESTION_TYPES = {
    "ice_breaker": [
        "Tell me about yourself.",
        "How did you hear about this position?",
        "Why are you interested in working for our company?",
        "What are you looking for in your next role?",
        "What are your strengths and weaknesses?"
    ],
    "behavioral": [
        "Tell me about a challenging project you worked on.",
        "Describe a situation where you had to work under pressure.",
        "How do you handle conflicts with team members?",
        "Tell me about a time you failed and what you learned from it.",
        "Describe your leadership style with an example."
    ],
    "technical_general": [
        "What programming languages are you most comfortable with?",
        "Describe your experience with databases.",
        "How do you approach testing your code?",
        "Explain your development workflow.",
        "How do you keep up with industry trends and new technologies?"
    ]
}

def initialize_interview(resume_data, jd_data, interview_config):
    """Initialize an interview session"""
    # Create RAG engine
    rag_engine = RAGEngine()
    rag_engine.create_knowledge_base(resume_data, jd_data)
    
    # Determine question distribution based on technical focus
    technical_focus = interview_config.get("technical_focus", 50)
    
    num_questions = calculate_num_questions(interview_config.get("duration", 20))
    
    # Initialize question plan
    question_plan = {
        "ice_breaker": 2,  # Always start with 2 ice breakers
        "behavioral": max(1, int((num_questions - 2) * (100 - technical_focus) / 100)),
        "technical": max(1, int((num_questions - 2) * technical_focus / 100)),
    }
    
    # Create interview session object
    interview_session = {
        "resume_data": resume_data,
        "jd_data": jd_data,
        "config": interview_config,
        "question_plan": question_plan,
        "current_stage": "ice_breaker",
        "questions_asked": 0,
        "max_questions": num_questions,
        "rag_engine": rag_engine
    }
    
    return interview_session

def calculate_num_questions(duration_minutes):
    """Calculate number of questions based on interview duration"""
    # Assume average of 2.5 minutes per question including answer time
    return max(5, int(duration_minutes / 2.5))

def get_next_question(interview_session):
    """Get the next interview question"""
    # Check if we've reached the maximum number of questions
    if interview_session["questions_asked"] >= interview_session["max_questions"]:
        return None
    
    # Determine current stage
    current_stage = interview_session["current_stage"]
    question_plan = interview_session["question_plan"]
    
    # If we've asked all questions in the current stage, move to the next stage
    if current_stage == "ice_breaker" and interview_session["questions_asked"] >= question_plan["ice_breaker"]:
        current_stage = "behavioral"
        interview_session["current_stage"] = current_stage
    elif current_stage == "behavioral" and interview_session["questions_asked"] >= (question_plan["ice_breaker"] + question_plan["behavioral"]):
        current_stage = "technical"
        interview_session["current_stage"] = current_stage
    
    # Generate question based on stage
    question = ""
    if current_stage == "ice_breaker":
        # Use predefined ice breaker questions
        question = random.choice(QUESTION_TYPES["ice_breaker"])
    elif current_stage == "behavioral":
        # Use predefined behavioral questions or RAG
        if random.random() < 0.3:  # 30% chance of using predefined
            question = random.choice(QUESTION_TYPES["behavioral"])
        else:
            # Use RAG to generate a more targeted behavioral question
            question = interview_session["rag_engine"].generate_question("behavioral", "middle")
    elif current_stage == "technical":
        # Use RAG to generate technical questions
        question = interview_session["rag_engine"].generate_question("technical", "late")
    
    # Increment questions asked
    interview_session["questions_asked"] += 1
    
    return question


# app/interview/personas.py
import random

# Define interviewer personas
PERSONAS = {
    "American (US)": {
        "Male": {
            "name": "Michael Johnson",
            "intro": "Hi there! I'm Michael Johnson. I'll be conducting your interview today. Let's have a productive conversation.",
            "speech_config": {
                "accent": "en-US",
                "voice_id": "en-US-Neural2-D",
                "speaking_rate": 1.0
            }
        },
        "Female": {
            "name": "Sarah Williams",
            "intro": "Hello! I'm Sarah Williams. I'm excited to learn more about your qualifications today.",
            "speech_config": {
                "accent": "en-US",
                "voice_id": "en-US-Neural2-F",
                "speaking_rate": 1.0
            }
        }
    },
    "British (UK)": {
        "Male": {
            "name": "James Smith",
            "intro": "Good day! I'm James Smith. Delighted to be speaking with you about this position.",
            "speech_config": {
                "accent": "en-GB",
                "voice_id": "en-GB-Neural2-B",
                "speaking_rate": 0.95
            }
        },
        "Female": {
            "name": "Emma Clarke",
            "intro": "Hello there! I'm Emma Clarke. I'll be conducting your interview today. Let's begin, shall we?",
            "speech_config": {
                "accent": "en-GB",
                "voice_id": "en-GB-Neural2-C",
                "speaking_rate": 0.95
            }
        }
    },
    "Australian": {
        "Male": {
            "name": "David Wilson",
            "intro": "G'day! I'm David Wilson. Looking forward to our chat about your experience.",
            "speech_config": {
                "accent": "en-AU",
                "voice_id": "en-AU-Neural2-B",
                "speaking_rate": 1.05
            }
        },
        "Female": {
            "name": "Olivia Taylor",
            "intro": "Hello! I'm Olivia Taylor. Thanks for joining me today for this interview.",
            "speech_config": {
                "accent": "en-AU",
                "voice_id": "en-AU-Neural2-A",
                "speaking_rate": 1.05
            }
        }
    },
    "Indian": {
        "Male": {
            "name": "Raj Patel",
            "intro": "Hello and welcome! I'm Raj Patel. I'm pleased to be interviewing you today.",
            "speech_config": {
                "accent": "en-IN",
                "voice_id": "en-IN-Neural2-B",
                "speaking_rate": 0.95
            }
        },
        "Female": {
            "name": "Priya Sharma",
            "intro": "Namaste! I'm Priya Sharma. I'll be conducting your interview today. Let's begin.",
            "speech_config": {
                "accent": "en-IN",
                "voice_id": "en-IN-Neural2-A",
                "speaking_rate": 0.95
            }
        }
    },
    "Chinese": {
        "Male": {
            "name": "Wei Zhang",
            "intro": "Hello! I'm Wei Zhang. Thank you for your interest in our company. Shall we begin the interview?",
            "speech_config": {
                "accent": "cmn-CN",
                "voice_id": "cmn-CN-Neural2-B",
                "speaking_rate": 1.0
            }
        },
        "Female": {
            "name": "Li Chen",
            "intro": "Hello! I'm Li Chen. I'm looking forward to discussing your qualifications today.",
            "speech_config": {
                "accent": "cmn-CN",
                "voice_id": "cmn-CN-Neural2-A",
                "speaking_rate": 1.0
            }
        }
    }
}

def get_interviewer_persona(accent, gender):
    """Get an interviewer persona based on accent and gender"""
    if accent in PERSONAS and gender in PERSONAS[accent]:
        return PERSONAS[accent][gender]
    
    # Fallback to random selection
    random_accent = random.choice(list(PERSONAS.keys()))
    random_gender = random.choice(["Male", "Female"])
    return PERSONAS[random_accent][random_gender]


# app/interview/session_recorder.py
import time
import json
import os
from datetime import datetime

class SessionRecorder:
    def __init__(self, user_id, session_id=None):
        """Initialize the session recorder"""
        self.user_id = user_id
        self.session_id = session_id or f"{int(time.time())}"
        self.start_time = datetime.now()
        self.end_time = None
        self.questions = []
        self.answers = []
        self.session_metadata = {}
        self.recording_path = None
    
    def add_question(self, question, question_type):
        """Add a question to the session"""
        question_data = {
            "text": question,
            "type": question_type,
            "timestamp": datetime.now().isoformat()
        }
        self.questions.append(question_data)
    
    def add_answer(self, answer_text, audio_segment=None):
        """Add an answer to the session"""
        answer_data = {
            "text": answer_text,
            "timestamp": datetime.now().isoformat()
        }
        
        if audio_segment:
            # In a real app, save the audio segment to storage
            # and reference it here
            answer_data["audio_reference"] = f"user_{self.user_id}_session_{self.session_id}_answer_{len(self.answers)}.wav"
        
        self.answers.append(answer_data)
    
    def set_recording_path(self, path):
        """Set the path for the recorded interview video"""
        self.recording_path = path
    
    def set_metadata(self, metadata):
        """Set session metadata"""
        self.session_metadata = metadata
    
    def end_session(self):
        """End the recording session"""
        self.end_time = datetime.now()
    
    def save_session_data(self, output_dir="session_recordings"):
        """Save the session data to a file"""
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        session_data = {
            "user_id": self.user_id,
            "session_id": self.session_id,
            "start_time": self.start_time.isoformat(),
            "end_time": (self.end_time or datetime.now()).isoformat(),
            "duration_seconds": (self.end_time or datetime.now() - self.start_time).total_seconds(),
            "questions": self.questions,
            "answers": self.answers,
            "metadata": self.session_metadata,
            "recording_path": self.recording_path
        }
        
        output_file = os.path.join(output_dir, f"session_{self.session_id}.json")
        
        with open(output_file, "w") as f:
            json.dump(session_data, f, indent=2)
        
        return output_file
    
    def get_session_summary(self):
        """Get a summary of the session"""
        return {
            "session_id": self.session_id,
            "start_time": self.start_time.isoformat(),
            "end_time": (self.end_time or datetime.now()).isoformat(),
            "duration_seconds": (self.end_time or datetime.now() - self.start_time).total_seconds(),
            "num_questions": len(self.questions),
            "metadata": self.session_metadata
        }


# app/livekit_integration.py
import os
import requests
import json
import base64
from datetime import datetime, timedelta
import jwt

class LiveKitManager:
    def __init__(self, api_key=None, api_secret=None, livekit_url=None):
        """Initialize LiveKit manager"""
        self.api_key = api_key or os.environ.get("LIVEKIT_API_KEY")
        self.api_secret = api_secret or os.environ.get("LIVEKIT_API_SECRET")
        self.livekit_url = livekit_url or os.environ.get("LIVEKIT_URL", "https://your-livekit-instance.livekit.cloud")
        
        if not self.api_key or not self.api_secret:
            raise ValueError("LiveKit API key and secret are required")
    
    def create_room(self, room_name=None):
        """Create a LiveKit room"""
        if not room_name:
            room_name = f"interview-{int(datetime.now().timestamp())}"
        
        # Generate JWT token for admin access
        token = self._generate_admin_token(room_name)
        
        # Create room via LiveKit API
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        }
        
        data = {
            "name": room_name,
            "empty_timeout": 5 * 60,  # 5 minutes
            "max_participants": 2
        }
        
        try:
            response = requests.post(
                f"{self.livekit_url}/twirp/livekit.RoomService/CreateRoom",
                headers=headers,
                data=json.dumps(data)
            )
            response.raise_for_status()
            return response.json()
        except Exception as e:
            print(f"Error creating LiveKit room: {e}")
            return None
    
    def generate_participant_token(self, room_name, participant_identity, participant_name, is_interviewer=False):
        """Generate a participant token for joining a room"""
        # Set token expiration (1 hour)
        exp = datetime.now() + timedelta(hours=1)
        
        # Create claims
        claims = {
            "video": {
                "room": room_name,
                "roomJoin": True,
                "canPublish": True,
                "canSubscribe": True,
                "canPublishData": True
            },
            "iss": self.api_key,
            "sub": participant_identity,
            "name": participant_name,
            "exp": int(exp.timestamp())
        }
        
        # Add interviewer-specific permissions if needed
        if is_interviewer:
            claims["video"]["roomAdmin"] = True
            claims["video"]["roomCreate"] = True
        
        # Generate token
        token = jwt.encode(claims, self.api_secret, algorithm="HS256")
        
        return token
    
    def _generate_admin_token(self, room_name=None):
        """Generate admin token for LiveKit API access"""
        # Set token expiration (5 minutes)
        exp = datetime.now() + timedelta(minutes=5)
        
        # Create claims
        claims = {
            "video": {
                "roomCreate": True,
                "roomList": True,
                "roomRecord": True,
                "roomAdmin": True
            },
            "iss": self.api_key,
            "sub": "admin",
            "exp": int(exp.timestamp())
        }
        
        # Add room-specific permissions if room name is provided
        if room_name:
            claims["video"]["room"] = room_name
        
        # Generate token
        token = jwt.encode(claims, self.api_secret, algorithm="HS256")
        
        return token
    
    def start_recording(self, room_name):
        """Start recording a LiveKit room"""
        # Generate JWT token for admin access
        token = self._generate_admin_token(room_name)
        
        # Start recording via LiveKit API
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        }
        
        data = {
            "room_name": room_name,
            "output_type": "mp4"
        }
        
        try:
            response = requests.post(
                f"{self.livekit_url}/twirp/livekit.Egress/StartRoomCompositeEgress",
                headers=headers,
                data=json.dumps(data)
            )
            response.raise_for_status()
            return response.json()
        except Exception as e:
            print(f"Error starting recording: {e}")
            return None


# app/auth/firebase_auth.py
import pyrebase
import json
import os

class FirebaseAuth:
    def __init__(self, config=None):
        """Initialize Firebase Authentication"""
        self.config = config or {
            "apiKey": os.environ.get("FIREBASE_API_KEY"),
            "authDomain": os.environ.get("FIREBASE_AUTH_DOMAIN"),
            "databaseURL": os.environ.get("FIREBASE_DATABASE_URL"),
            "projectId": os.environ.get("FIREBASE_PROJECT_ID"),
            "storageBucket": os.environ.get("FIREBASE_STORAGE_BUCKET"),
            "messagingSenderId": os.environ.get("FIREBASE_MESSAGING_SENDER_ID"),
            "appId": os.environ.get("FIREBASE_APP_ID")
        }
        
        self.firebase = pyrebase.initialize_app(self.config)
        self.auth = self.firebase.auth()
        self.db = self.firebase.database()
    
    def sign_up(self, email, password, display_name):
        """Register a new user"""
        try:
            # Create user in Firebase Authentication
            user = self.auth.create_user_with_email_and_password(email, password)
            
            # Add user profile to Realtime Database
            user_id = user['localId']
            self.db.child("users").child(user_id).set({
                "email": email,
                "display_name": display_name,
                "created_at": {".sv": "timestamp"}
            })
            
            return True, user
        except Exception as e:
            print(f"Error signing up: {e}")
            return False, str(e)
    
    def sign_in(self, email, password):
        """Sign in an existing user"""
        try:
            user = self.auth.sign_in_with_email_and_password(email, password)
            return True, user
        except Exception as e:
            print(f"Error signing in: {e}")
            return False, str(e)
    
    def get_user_data(self, user_id, token):
        """Get user data from Realtime Database"""
        try:
            user_data = self.db.child("users").child(user_id).get(token=token)
            return user_data.val()
        except Exception as e:
            print(f"Error getting user data: {e}")
            return None
    
    def reset_password(self, email):
        """Send password reset email"""
        try:
            self.auth.send_password_reset_email(email)
            return True, "Password reset email sent"
        except Exception as e:
            print(f"Error sending password reset: {e}")
            return False, str(e)


# app/database/firebase_storage.py
import pyrebase
import os
import uuid
from datetime import datetime

class FirebaseStorage:
    def __init__(self, config=None):
        """Initialize Firebase Storage"""
        self.config = config or {
            "apiKey": os.environ.get("FIREBASE_API_KEY"),
            "authDomain": os.environ.get("FIREBASE_AUTH_DOMAIN"),
            "databaseURL": os.environ.get("FIREBASE_DATABASE_URL"),
            "projectId": os.environ.get("FIREBASE_PROJECT_ID"),
            "storageBucket": os.environ.get("FIREBASE_STORAGE_BUCKET"),
            "messagingSenderId": os.environ.get("FIREBASE_MESSAGING_SENDER_ID"),
            "appId": os.environ.get("FIREBASE_APP_ID")
        }
        
        self.firebase = pyrebase.initialize_app(self.config)
        self.storage = self.firebase.storage()
    
    def upload_file(self, file_path, folder, token=None):
        """Upload a file to Firebase Storage"""
        try:
            # Generate unique filename
            file_name = os.path.basename(file_path)
            timestamp = int(datetime.now().timestamp())
            unique_name = f"{timestamp}_{file_name}"
            
            # Upload file
            storage_path = f"{folder}/{unique_name}"
            self.storage.child(storage_path).put(file_path, token=token)
            
            # Get download URL
            url = self.storage.child(storage_path).get_url(token=token)
            
            return True, url
        except Exception as e:
            print(f"Error uploading file: {e}")
            return False, str(e)
    
    def upload_file_from_memory(self, file_object, file_name, folder, token=None):
        """Upload a file object to Firebase Storage"""
        try:
            # Generate unique filename
            timestamp = int(datetime.now().timestamp())
            unique_name = f"{timestamp}_{file_name}"
            
            # Upload file
            storage_path = f"{folder}/{unique_name}"
            
            # Save temp file first
            temp_path = f"/tmp/{unique_name}"
            with open(temp_path, "wb") as f:
                f.write(file_object.getvalue())
            
            # Upload to Firebase
            self.storage.child(storage_path).put(temp_path, token=token)
            
            # Clean up temp file
            os.remove(temp_path)
            
            # Get download URL
            url = self.storage.child(storage_path).get_url(token=token)
            
            return True, url
        except Exception as e:
            print(f"Error uploading file: {e}")
            return False, str(e)
    
    def delete_file(self, file_path, token=None):
        """Delete a file from Firebase Storage"""
        try:
            self.storage.delete(file_path, token=token)
            return True, "File deleted successfully"
        except Exception as e:
            print(f"Error deleting file: {e}")
            return False, str(e)
