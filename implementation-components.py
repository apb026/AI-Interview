# app/document_processing/resume_parser.py
import io
import re
import PyPDF2
import docx
from pdfminer.high_level import extract_text

def parse_resume(uploaded_file):
    """
    Parse resume file (PDF, DOCX, TXT) and extract relevant information
    """
    file_type = uploaded_file.name.split('.')[-1].lower()
    content = None
    
    # Extract text based on file type
    if file_type == 'pdf':
        content = extract_text_from_pdf(uploaded_file)
    elif file_type == 'docx':
        content = extract_text_from_docx(uploaded_file)
    elif file_type == 'txt':
        content = uploaded_file.getvalue().decode('utf-8')
    else:
        return None
    
    # Parse the content to extract structured information
    parsed_data = extract_resume_information(content)
    return parsed_data

def extract_text_from_pdf(pdf_file):
    """Extract text content from PDF file"""
    text = ""
    try:
        # First approach with PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_file.getvalue()))
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        # If PyPDF2 didn't extract text well, try pdfminer
        if not text.strip():
            text = extract_text(io.BytesIO(pdf_file.getvalue()))
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        # Fallback to pdfminer
        try:
            text = extract_text(io.BytesIO(pdf_file.getvalue()))
        except Exception as e2:
            print(f"Error with fallback PDF extraction: {e2}")
    
    return text

def extract_text_from_docx(docx_file):
    """Extract text content from DOCX file"""
    try:
        doc = docx.Document(io.BytesIO(docx_file.getvalue()))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

def extract_resume_information(text):
    """
    Extract structured information from the resume text
    """
    # Basic extraction - in a real app, use more sophisticated NLP
    resume_data = {
        'raw_text': text,
        'name': extract_name(text),
        'email': extract_email(text),
        'phone': extract_phone(text),
        'education': extract_education(text),
        'experience': extract_experience(text),
        'skills': extract_skills(text)
    }
    
    return resume_data

def extract_name(text):
    # Simple name extraction - first line of the resume often contains the name
    lines = text.split('\n')
    for line in lines[:5]:  # Check first 5 lines
        if line.strip() and not re.search(r'@|resume|cv|curriculum', line.lower()):
            return line.strip()
    return ""

def extract_email(text):
    # Email extraction using regex
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    emails = re.findall(email_pattern, text)
    return emails[0] if emails else ""

def extract_phone(text):
    # Phone number extraction using regex
    phone_pattern = r'(\+\d{1,2}\s?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}'
    phones = re.findall(phone_pattern, text)
    return phones[0] if phones else ""

def extract_education(text):
    # Basic education extraction
    edu_keywords = ['education', 'university', 'college', 'bachelor', 'master', 'phd', 'degree']
    education_section = extract_section(text, edu_keywords)
    
    # In a real app, parse this section further to extract institutions, degrees, years
    return education_section

def extract_experience(text):
    # Basic work experience extraction
    exp_keywords = ['experience', 'employment', 'work history', 'professional experience']
    experience_section = extract_section(text, exp_keywords)
    
    # In a real app, parse this further to extract companies, roles, dates, responsibilities
    return experience_section

def extract_skills(text):
    # Basic skills extraction
    skill_keywords = ['skills', 'technical skills', 'competencies', 'proficiencies']
    skills_section = extract_section(text, skill_keywords)
    
    # Extract individual skills
    # In a real app, use a predefined list of skills to match against
    return skills_section

def extract_section(text, keywords):
    """Extract a section from the text based on keywords"""
    text_lower = text.lower()
    for keyword in keywords:
        if keyword in text_lower:
            # Find the starting position of the keyword
            start_pos = text_lower.find(keyword)
            
            # Extract from this position to the next section
            section_text = text[start_pos:start_pos + 1000]  # Limit to 1000 chars for simplicity
            
            # In a real app, find the end of the section more intelligently
            return section_text
    
    return ""


# app/document_processing/jd_parser.py
import io
import re
import PyPDF2
import docx
from pdfminer.high_level import extract_text

def parse_jd_file(uploaded_file):
    """Parse job description from uploaded file"""
    file_type = uploaded_file.name.split('.')[-1].lower()
    content = None
    
    # Extract text based on file type
    if file_type == 'pdf':
        content = extract_text_from_pdf(uploaded_file)
    elif file_type == 'docx':
        content = extract_text_from_docx(uploaded_file)
    elif file_type == 'txt':
        content = uploaded_file.getvalue().decode('utf-8')
    else:
        return None
    
    # Parse the content to extract structured information
    return parse_jd_text(content)

def extract_text_from_pdf(pdf_file):
    """Extract text content from PDF file"""
    # Same implementation as in resume_parser.py
    text = ""
    try:
        # First approach with PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_file.getvalue()))
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        # If PyPDF2 didn't extract text well, try pdfminer
        if not text.strip():
            text = extract_text(io.BytesIO(pdf_file.getvalue()))
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        # Fallback to pdfminer
        try:
            text = extract_text(io.BytesIO(pdf_file.getvalue()))
        except Exception as e2:
            print(f"Error with fallback PDF extraction: {e2}")
    
    return text

def extract_text_from_docx(docx_file):
    """Extract text content from DOCX file"""
    # Same implementation as in resume_parser.py
    try:
        doc = docx.Document(io.BytesIO(docx_file.getvalue()))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

def parse_jd_text(text):
    """Parse job description from text"""
    jd_data = {
        'raw_text': text,
        'title': extract_job_title(text),
        'company': extract_company_name(text),
        'required_skills': extract_required_skills(text),
        'responsibilities': extract_responsibilities(text),
        'qualifications': extract_qualifications(text),
        'location': extract_location(text)
    }
    
    return jd_data

def extract_job_title(text):
    """Extract job title from JD text"""
    # Simple approach - look for common patterns
    title_patterns = [
        r'job title\s*[:-]\s*(.*?)[\n\r]',
        r'position\s*[:-]\s*(.*?)[\n\r]',
        r'role\s*[:-]\s*(.*?)[\n\r]'
    ]
    
    for pattern in title_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    
    # Fallback: First line is often the job title
    first_line = text.strip().split('\n')[0]
    if len(first_line) < 100:  # Reasonable length for a title
        return first_line
    
    return "Unknown Position"

def extract_company_name(text):
    """Extract company name from JD text"""
    company_patterns = [
        r'company\s*[:-]\s*(.*?)[\n\r]',
        r'about\s+([a-z0-9\s]+)[\n\r]',
        r'at\s+([a-z0-9\s]+)[\n\r]',
        r'join\s+([a-z0-9\s]+)[\n\r]'
    ]
    
    for pattern in company_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    
    return "Unknown Company"

def extract_required_skills(text):
    """Extract required skills from JD text"""
    # Look for skills section
    skills_section = extract_section(text, ['skills', 'requirements', 'qualifications'])
    
    # In a real app, use more sophisticated NLP to extract specific skills
    return skills_section

def extract_responsibilities(text):
    """Extract job responsibilities from JD text"""
    # Look for responsibilities section
    resp_section = extract_section(text, ['responsibilities', 'duties', 'what you\'ll do'])
    
    return resp_section

def extract_qualifications(text):
    """Extract required qualifications from JD text"""
    # Look for qualifications section
    qual_section = extract_section(text, ['qualifications', 'requirements', 'what we\'re looking for'])
    
    return qual_section

def extract_location(text):
    """Extract job location from JD text"""
    location_patterns = [
        r'location\s*[:-]\s*(.*?)[\n\r]',
        r'based in\s+(.*?)[\n\r]',
        r'position is (?:located|based) in\s+(.*?)[\n\r]'
    ]
    
    for pattern in location_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    
    return "Unknown Location"

def extract_section(text, keywords):
    """Extract a section from the text based on keywords"""
    text_lower = text.lower()
    for keyword in keywords:
        if keyword in text_lower:
            # Find the starting position of the keyword
            start_pos = text_lower.find(keyword)
            
            # Extract from this position to the next section
            section_text = text[start_pos:start_pos + 1000]  # Limit to 1000 chars for simplicity
            
            # In a real app, find the end of the section more intelligently
            return section_text
    
    return ""


# app/document_processing/web_scraper.py
import requests
from bs4 import BeautifulSoup
import re

def scrape_job_description(url):
    """Scrape job description from a URL"""
    try:
        # Get the webpage content
        response = requests.get(url, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
        response.raise_for_status()
        
        # Parse HTML
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Extract job information based on common patterns
        # This is simplified - in a real app, you would need site-specific extractors
        
        # Extract main content - this is highly site-dependent
        main_content = extract_main_content(soup)
        
        # Parse the extracted content
        from app.document_processing.jd_parser import parse_jd_text
        jd_data = parse_jd_text(main_content)
        
        # Add source URL
        jd_data['source_url'] = url
        
        return jd_data
    
    except Exception as e:
        print(f"Error scraping job description: {e}")
        return None

def extract_main_content(soup):
    """Extract the main content from the HTML"""
    # Try different common containers
    content = ""
    
    # Try article tag
    article = soup.find('article')
    if article:
        content = article.get_text()
    
    # Try job description specific divs
    if not content:
        job_divs = soup.find_all(['div', 'section'], class_=re.compile(r'job(-|\s)?desc|description|posting|details', re.I))
        if job_divs:
            content = job_divs[0].get_text()
    
    # Fallback to main content area
    if not content:
        main = soup.find(['main', 'div'], id=re.compile(r'main|content', re.I))
        if main:
            content = main.get_text()
    
    # Last resort - get all text
    if not content:
        content = soup.get_text()
    
    # Clean up the text
    content = re.sub(r'\s+', ' ', content).strip()
    
    return content


# app/interview/rag_engine.py
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.text_splitter import Rec